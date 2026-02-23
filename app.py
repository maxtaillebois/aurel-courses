import streamlit as st
import json
import os
import io
import requests
from datetime import datetime
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Configuration ---
st.set_page_config(page_title="🛒 Liste de courses", page_icon="🛒", layout="wide")

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
load_dotenv(os.path.join(BASE_DIR, ".env"))

RECETTES_PATH = os.path.join(BASE_DIR, "recettes.json")
CATALOGUE_PATH = os.path.join(BASE_DIR, "catalogue.json")
NOTION_TOKEN = os.getenv("NOTION_TOKEN")
NOTION_PAGE_ID = os.getenv("NOTION_PAGE_ID")

UNITES = ["pièce", "g", "kg", "ml", "cl", "L"]


# --- Chargement / sauvegarde ---
def load_recettes():
    with open(RECETTES_PATH, "r", encoding="utf-8") as f:
        return json.load(f)["plats"]


def load_catalogue():
    with open(CATALOGUE_PATH, "r", encoding="utf-8") as f:
        return json.load(f)["rayons"]


def save_recettes(plats):
    with open(RECETTES_PATH, "w", encoding="utf-8") as f:
        json.dump({"plats": plats}, f, ensure_ascii=False, indent=2)


def save_catalogue(rayons):
    with open(CATALOGUE_PATH, "w", encoding="utf-8") as f:
        json.dump({"rayons": rayons}, f, ensure_ascii=False, indent=2)


# --- Utilitaires ---
def format_item(nom, quantite, unite):
    """Formate un article pour l'affichage : 'Carottes — 450g' ou 'Poulet — x1'."""
    if unite == "pièce":
        if quantite == 1:
            return nom
        return f"{nom} (x{quantite})"
    return f"{nom} ({quantite}{unite})"


def merge_ingredients(ingredients_list):
    """Fusionne les ingrédients structurés en dédoublonnant et cumulant les quantités.
    Entrée: liste de {"nom", "rayon", "quantite", "unite"}
    Retourne: {rayon: [(nom, quantite, unite), ...]}
    """
    merged = {}
    for ing in ingredients_list:
        nom = ing["nom"]
        rayon = ing["rayon"]
        qty = ing.get("quantite", 1)
        unite = ing.get("unite", "pièce")
        key = (nom.lower(), rayon)

        if key in merged:
            if merged[key]["unite"] == unite:
                merged[key]["quantite"] += qty
            else:
                # Unités différentes → garder séparés (ajouter comme nouvel item)
                alt_key = (nom.lower() + f"_{unite}", rayon)
                if alt_key in merged:
                    merged[alt_key]["quantite"] += qty
                else:
                    merged[alt_key] = {
                        "rayon": rayon,
                        "nom": nom,
                        "quantite": qty,
                        "unite": unite,
                    }
        else:
            merged[key] = {
                "rayon": rayon,
                "nom": nom,
                "quantite": qty,
                "unite": unite,
            }

    result = {}
    for data in merged.values():
        rayon = data["rayon"]
        if rayon not in result:
            result[rayon] = []
        result[rayon].append((data["nom"], data["quantite"], data["unite"]))

    # Trier par nom au sein de chaque rayon
    for rayon in result:
        result[rayon].sort(key=lambda x: x[0].lower())

    return result


def get_recipe_ingredients(recettes, selected_names):
    """Récupère tous les ingrédients des recettes sélectionnées."""
    ingredients = []
    for recette in recettes:
        if recette["nom"] in selected_names:
            ingredients.extend(recette["ingredients"])
    return ingredients


def build_final_list(recipe_items_by_rayon, free_items_by_rayon):
    """Combine recettes + produits libres, par rayon.
    Les deux entrées sont {rayon: [(nom, quantite, unite), ...]}
    Retourne le même format, fusionné et trié par rayon.
    """
    # Rassembler tous les items dans une seule liste pour fusion
    all_items = []
    for rayon, items in recipe_items_by_rayon.items():
        for nom, qty, unite in items:
            all_items.append({"nom": nom, "rayon": rayon, "quantite": qty, "unite": unite})
    for rayon, items in free_items_by_rayon.items():
        for nom, qty, unite in items:
            all_items.append({"nom": nom, "rayon": rayon, "quantite": qty, "unite": unite})

    merged = merge_ingredients(all_items)

    # Ordonner par rayon
    rayon_order = [
        "BOULANGERIE", "LÉGUMES", "FRUITS", "AIL & FINES HERBES",
        "CHARCUTERIE", "TRAITEUR", "POISSONNERIE", "BOUCHERIE",
        "SURGELÉS", "FROMAGES", "YAOURTS", "PRODUITS LAITIERS",
        "ÉPICERIE SALÉE", "CUISINE DU MONDE", "ÉPICERIE SUCRÉE",
        "BOISSONS", "NOURRITURE BÉBÉ", "HYGIÈNE & DIVERS",
    ]

    final = {}
    for rayon in rayon_order:
        if rayon in merged:
            final[rayon] = merged[rayon]

    for rayon in sorted(set(merged.keys()) - set(rayon_order)):
        final[rayon] = merged[rayon]

    return final


def subtract_stock(final_list, stock_items):
    """Soustrait le stock de la liste finale.
    stock_items: {rayon: [(nom, quantite, unite), ...]}
    Retourne la liste finale nettoyée.
    """
    # Indexer le stock par (nom.lower(), rayon)
    stock_index = {}
    for rayon, items in stock_items.items():
        for nom, qty, unite in items:
            key = (nom.lower(), rayon)
            if key in stock_index:
                if stock_index[key]["unite"] == unite:
                    stock_index[key]["quantite"] += qty
                else:
                    stock_index[key] = {"quantite": qty, "unite": unite}
            else:
                stock_index[key] = {"quantite": qty, "unite": unite}

    result = {}
    for rayon, items in final_list.items():
        new_items = []
        for nom, qty, unite in items:
            key = (nom.lower(), rayon)
            if key in stock_index:
                stock = stock_index[key]
                if stock["unite"] == unite:
                    remaining = qty - stock["quantite"]
                    if remaining > 0:
                        new_items.append((nom, remaining, unite))
                    # Si remaining <= 0, on ne l'ajoute pas (couvert par le stock)
                else:
                    # Unités différentes → on ne peut pas soustraire, on garde tel quel
                    new_items.append((nom, qty, unite))
            else:
                new_items.append((nom, qty, unite))
        if new_items:
            result[rayon] = new_items

    return result


def add_ingredient_to_catalogue(catalogue, nom_ingredient, rayon_nom):
    """Ajoute un ingrédient au catalogue s'il n'y est pas déjà."""
    for rayon in catalogue:
        if rayon["nom"] == rayon_nom:
            existing_lower = [a.lower() for a in rayon["articles"]]
            if nom_ingredient.lower() not in existing_lower:
                rayon["articles"].append(nom_ingredient)
                rayon["articles"].sort(key=str.lower)
                return True
            return False
    catalogue.append({"nom": rayon_nom, "articles": [nom_ingredient]})
    return True


def export_to_notion(final_list, selected_recipes):
    """Crée une page Notion avec des cases à cocher via l'API."""
    if not NOTION_TOKEN or not NOTION_PAGE_ID:
        return False, "Configuration Notion manquante. Vérifiez le fichier .env.", None

    headers = {
        "Authorization": f"Bearer {NOTION_TOKEN}",
        "Content-Type": "application/json",
        "Notion-Version": "2022-06-28",
    }

    date_str = datetime.now().strftime("%d/%m/%Y")
    title = f"🛒 Liste de courses — {date_str}"

    children = []

    if selected_recipes:
        children.append({
            "object": "block",
            "type": "paragraph",
            "paragraph": {
                "rich_text": [{
                    "type": "text",
                    "text": {"content": f"🍽️ {' • '.join(selected_recipes)}"},
                    "annotations": {"italic": True, "color": "gray"},
                }]
            }
        })
        children.append({"object": "block", "type": "divider", "divider": {}})

    for rayon, items in final_list.items():
        children.append({
            "object": "block",
            "type": "heading_2",
            "heading_2": {
                "rich_text": [{"type": "text", "text": {"content": rayon}}]
            }
        })
        for nom, qty, unite in items:
            display = format_item(nom, qty, unite)
            children.append({
                "object": "block",
                "type": "to_do",
                "to_do": {
                    "rich_text": [{"type": "text", "text": {"content": display}}],
                    "checked": False,
                }
            })

    payload = {
        "parent": {"page_id": NOTION_PAGE_ID},
        "properties": {
            "title": [{"text": {"content": title}}]
        },
        "children": children[:100],
    }

    try:
        resp = requests.post(
            "https://api.notion.com/v1/pages",
            headers=headers,
            json=payload,
            timeout=15,
        )

        if resp.status_code == 200:
            page_url = resp.json().get("url", "")

            if len(children) > 100:
                page_id = resp.json()["id"]
                for i in range(100, len(children), 100):
                    batch = children[i:i+100]
                    requests.patch(
                        f"https://api.notion.com/v1/blocks/{page_id}/children",
                        headers=headers,
                        json={"children": batch},
                        timeout=15,
                    )

            return True, "Page créée dans Notion !", page_url
        else:
            error = resp.json().get("message", resp.text)
            return False, f"Erreur Notion : {error}", None

    except requests.exceptions.Timeout:
        return False, "Timeout : Notion n'a pas répondu.", None
    except Exception as e:
        return False, f"Erreur : {str(e)}", None


def export_to_docx(final_list, selected_recipes):
    """Génère un fichier Word de la liste de courses."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    title = doc.add_heading("Liste de courses", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Semaine du {datetime.now().strftime('%d/%m/%Y')}")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(100, 100, 100)

    if selected_recipes:
        doc.add_paragraph()
        plats_para = doc.add_paragraph()
        plats_run = plats_para.add_run("Plats : ")
        plats_run.bold = True
        plats_run.font.size = Pt(10)
        plats_text = plats_para.add_run(" • ".join(selected_recipes))
        plats_text.font.size = Pt(10)
        plats_text.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()

    for rayon, items in final_list.items():
        heading = doc.add_heading(rayon, level=2)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(46, 117, 182)
            run.font.size = Pt(13)

        for nom, qty, unite in items:
            display = format_item(nom, qty, unite)
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(display)
            run.font.size = Pt(11)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# --- Chargement ---
recettes = load_recettes()
catalogue = load_catalogue()

# --- Session state ---
if "checked_items" not in st.session_state:
    st.session_state.checked_items = set()
if "new_recipe_ingredients" not in st.session_state:
    st.session_state.new_recipe_ingredients = []
if "editing_recipe" not in st.session_state:
    st.session_state.editing_recipe = None


# --- Interface ---
st.title("🛒 Liste de courses")

tab_recettes, tab_produits, tab_stock, tab_liste, tab_gerer = st.tabs(
    ["🍽️ Recettes", "🏪 Produits", "🏠 Mon stock", "📋 Ma liste", "✏️ Gérer les recettes"]
)

# =====================
# ONGLET 1 : RECETTES
# =====================
with tab_recettes:
    st.header("Sélectionnez vos plats de la semaine")

    search_recettes = st.text_input(
        "🔍 Rechercher une recette", key="search_recettes",
        placeholder="Ex : quiche, poulet...",
    )

    recettes_triees = sorted(recettes, key=lambda r: r["nom"].lower())
    if search_recettes.strip():
        q = search_recettes.strip().lower()
        recettes_triees = [r for r in recettes_triees if q in r["nom"].lower()]

    for recette in recettes_triees:
        ingredients_str = ", ".join(
            format_item(ing["nom"], ing.get("quantite", 1), ing.get("unite", "pièce"))
            for ing in recette["ingredients"]
        )
        st.checkbox(
            recette["nom"],
            key=f"recette_{recette['nom']}",
            help=ingredients_str,
        )

    _selected = [r["nom"] for r in recettes if st.session_state.get(f"recette_{r['nom']}", False)]
    if _selected:
        st.divider()
        st.subheader("Ingrédients sélectionnés")
        _ingredients = get_recipe_ingredients(recettes, _selected)
        _by_rayon = merge_ingredients(_ingredients)
        for rayon, items in sorted(_by_rayon.items()):
            st.markdown(f"**{rayon}**")
            for nom, qty, unite in sorted(items, key=lambda x: x[0].lower()):
                st.markdown(f"- {format_item(nom, qty, unite)}")

# =====================
# ONGLET 2 : PRODUITS
# =====================
with tab_produits:
    st.header("Ajoutez des articles par rayon")

    search_produits = st.text_input(
        "🔍 Rechercher un produit", key="search_produits",
        placeholder="Ex : yaourt, tomate...",
    )

    q_produits = search_produits.strip().lower() if search_produits.strip() else ""

    for rayon in catalogue:
        if q_produits:
            matching = [(j, a) for j, a in enumerate(rayon["articles"]) if q_produits in a.lower()]
            if not matching:
                continue
        else:
            matching = list(enumerate(rayon["articles"]))

        with st.expander(f"🏷️ {rayon['nom']} ({len(matching)} articles)", expanded=bool(q_produits)):
            for j, article in matching:
                cat_key = f"cat_{rayon['nom']}_{j}"
                qty_key = f"qty_{rayon['nom']}_{j}"

                col_check, col_qty = st.columns([3, 1])
                with col_check:
                    checked = st.checkbox(article, key=cat_key)
                with col_qty:
                    if checked:
                        st.number_input(
                            "Qté",
                            min_value=1,
                            value=st.session_state.get(qty_key, 1),
                            key=qty_key,
                            label_visibility="collapsed",
                        )

# =====================
# ONGLET 3 : MON STOCK
# =====================
with tab_stock:
    st.header("Ce que j'ai déjà")
    st.caption("Renseignez les produits que vous avez déjà chez vous. Ils seront soustraits de la liste finale.")

    search_stock = st.text_input(
        "🔍 Rechercher un produit", key="search_stock",
        placeholder="Ex : crème, riz...",
    )

    q_stock = search_stock.strip().lower() if search_stock.strip() else ""

    for rayon in catalogue:
        if q_stock:
            matching = [(j, a) for j, a in enumerate(rayon["articles"]) if q_stock in a.lower()]
            if not matching:
                continue
        else:
            matching = list(enumerate(rayon["articles"]))

        with st.expander(f"🏷️ {rayon['nom']} ({len(matching)} articles)", expanded=bool(q_stock)):
            for j, article in matching:
                stock_key = f"stock_{rayon['nom']}_{j}"
                stock_qty_key = f"stock_qty_{rayon['nom']}_{j}"
                stock_unit_key = f"stock_unit_{rayon['nom']}_{j}"

                col_check, col_qty, col_unit = st.columns([3, 1, 1])
                with col_check:
                    checked = st.checkbox(article, key=stock_key)
                with col_qty:
                    if checked:
                        st.number_input(
                            "Qté",
                            min_value=1,
                            value=st.session_state.get(stock_qty_key, 1),
                            key=stock_qty_key,
                            label_visibility="collapsed",
                        )
                with col_unit:
                    if checked:
                        st.selectbox(
                            "Unité",
                            options=UNITES,
                            index=0,
                            key=stock_unit_key,
                            label_visibility="collapsed",
                        )

# ==============================
# ONGLET 4 : GÉRER LES RECETTES
# ==============================
with tab_gerer:
    st.header("Gérer les recettes")

    recettes_noms = sorted([r["nom"] for r in recettes], key=str.lower)
    options = ["-- Nouvelle recette --"] + recettes_noms

    select_index = 0
    if st.session_state.editing_recipe and st.session_state.editing_recipe in recettes_noms:
        select_index = options.index(st.session_state.editing_recipe)

    choix = st.selectbox(
        "Choisir une recette à modifier ou en créer une nouvelle",
        options=options,
        index=select_index,
        key="recipe_selector",
    )

    is_editing = choix != "-- Nouvelle recette --"

    if is_editing:
        if st.session_state.editing_recipe != choix:
            st.session_state.editing_recipe = choix
            for r in recettes:
                if r["nom"] == choix:
                    st.session_state.new_recipe_ingredients = [
                        {
                            "nom": ing["nom"],
                            "rayon": ing["rayon"],
                            "quantite": ing.get("quantite", 1),
                            "unite": ing.get("unite", "pièce"),
                        }
                        for ing in r["ingredients"]
                    ]
                    break
            st.rerun()
    else:
        if st.session_state.editing_recipe is not None:
            st.session_state.editing_recipe = None
            st.session_state.new_recipe_ingredients = []
            st.rerun()

    default_name = choix if is_editing else ""
    recipe_name = st.text_input(
        "Nom de la recette",
        value=default_name,
        placeholder="Ex : Gratin dauphinois",
        key="edit_recipe_name",
    )

    st.divider()
    st.subheader("Ingrédients")

    rayon_names = [r["nom"] for r in catalogue]

    with st.form("add_ingredient_form", clear_on_submit=True):
        col_ing, col_rayon, col_qty, col_unit = st.columns([3, 2, 1, 1])
        with col_ing:
            ing_name = st.text_input(
                "Nom de l'ingrédient",
                placeholder="Ex : Pommes de terre",
            )
        with col_rayon:
            ing_rayon = st.selectbox("Rayon", options=rayon_names, index=0)
        with col_qty:
            ing_qty = st.number_input("Qté", min_value=1, value=1)
        with col_unit:
            ing_unite = st.selectbox("Unité", options=UNITES, index=0)
        submitted = st.form_submit_button("➕ Ajouter l'ingrédient")

        if submitted and ing_name.strip():
            st.session_state.new_recipe_ingredients.append({
                "nom": ing_name.strip(),
                "rayon": ing_rayon,
                "quantite": ing_qty,
                "unite": ing_unite,
            })

    if st.session_state.new_recipe_ingredients:
        st.markdown("---")
        st.markdown("**Ingrédients :**")
        to_remove = None
        for idx, ing in enumerate(st.session_state.new_recipe_ingredients):
            col_display, col_del = st.columns([4, 1])
            with col_display:
                display = format_item(ing["nom"], ing.get("quantite", 1), ing.get("unite", "pièce"))
                st.markdown(f"- **{display}** — _{ing['rayon']}_")
            with col_del:
                if st.button("🗑️", key=f"del_ing_{idx}"):
                    to_remove = idx

        if to_remove is not None:
            st.session_state.new_recipe_ingredients.pop(to_remove)
            st.rerun()

        st.markdown("---")

        if is_editing:
            col_save, col_delete = st.columns(2)
            with col_save:
                if st.button("💾 Enregistrer les modifications", type="primary"):
                    if not recipe_name.strip():
                        st.error("Donnez un nom à la recette.")
                    else:
                        new_name = recipe_name.strip()
                        existing_names = [r["nom"].lower() for r in recettes if r["nom"] != choix]
                        if new_name.lower() in existing_names:
                            st.error(f"La recette « {new_name} » existe déjà.")
                        else:
                            for r in recettes:
                                if r["nom"] == choix:
                                    r["nom"] = new_name
                                    r["ingredients"] = list(st.session_state.new_recipe_ingredients)
                                    break
                            save_recettes(recettes)

                            catalogue_modified = False
                            for ing in st.session_state.new_recipe_ingredients:
                                if add_ingredient_to_catalogue(catalogue, ing["nom"], ing["rayon"]):
                                    catalogue_modified = True
                            if catalogue_modified:
                                save_catalogue(catalogue)

                            st.session_state.editing_recipe = new_name
                            st.session_state.new_recipe_ingredients = []
                            st.success(f"✅ Recette « {new_name} » mise à jour !")
                            st.rerun()

            with col_delete:
                if st.button("🗑️ Supprimer cette recette", type="secondary"):
                    st.session_state["confirm_delete"] = True

                if st.session_state.get("confirm_delete", False):
                    st.warning(f"⚠️ Supprimer définitivement « {choix} » ?")
                    col_yes, col_no = st.columns(2)
                    with col_yes:
                        if st.button("Oui, supprimer", type="primary"):
                            recettes[:] = [r for r in recettes if r["nom"] != choix]
                            save_recettes(recettes)
                            st.session_state.editing_recipe = None
                            st.session_state.new_recipe_ingredients = []
                            st.session_state["confirm_delete"] = False
                            st.success(f"🗑️ Recette « {choix} » supprimée.")
                            st.rerun()
                    with col_no:
                        if st.button("Annuler"):
                            st.session_state["confirm_delete"] = False
                            st.rerun()
        else:
            if st.button("💾 Enregistrer la recette", type="primary"):
                if not recipe_name.strip():
                    st.error("Donnez un nom à la recette.")
                else:
                    existing_names = [r["nom"].lower() for r in recettes]
                    if recipe_name.strip().lower() in existing_names:
                        st.error(f"La recette « {recipe_name.strip()} » existe déjà.")
                    else:
                        new_recipe = {
                            "nom": recipe_name.strip(),
                            "ingredients": list(st.session_state.new_recipe_ingredients),
                        }
                        recettes.append(new_recipe)
                        save_recettes(recettes)

                        catalogue_modified = False
                        for ing in st.session_state.new_recipe_ingredients:
                            if add_ingredient_to_catalogue(catalogue, ing["nom"], ing["rayon"]):
                                catalogue_modified = True
                        if catalogue_modified:
                            save_catalogue(catalogue)

                        st.session_state.new_recipe_ingredients = []
                        st.success(f"✅ Recette « {recipe_name.strip()} » enregistrée !")
                        st.balloons()
                        st.rerun()
    else:
        st.caption("Ajoutez des ingrédients via le formulaire ci-dessus.")

# ============================================
# CALCUL DE LA LISTE FINALE (hors des tabs)
# ============================================
selected_recipes_final = []
for recette in recettes:
    if st.session_state.get(f"recette_{recette['nom']}", False):
        selected_recipes_final.append(recette["nom"])

recipe_ingredients_final = get_recipe_ingredients(recettes, selected_recipes_final)
recipe_by_rayon_final = merge_ingredients(recipe_ingredients_final)

# Produits cochés avec quantités
free_items_final = {}
for rayon in catalogue:
    items = []
    for j, article in enumerate(rayon["articles"]):
        cat_key = f"cat_{rayon['nom']}_{j}"
        qty_key = f"qty_{rayon['nom']}_{j}"
        if st.session_state.get(cat_key, False):
            qty = st.session_state.get(qty_key, 1)
            items.append((article, qty, "pièce"))
    if items:
        free_items_final[rayon["nom"]] = items

final_list_before_stock = build_final_list(recipe_by_rayon_final, free_items_final)

# Stock : produits déjà en possession
stock_items_final = {}
for rayon in catalogue:
    items = []
    for j, article in enumerate(rayon["articles"]):
        stock_key = f"stock_{rayon['nom']}_{j}"
        stock_qty_key = f"stock_qty_{rayon['nom']}_{j}"
        stock_unit_key = f"stock_unit_{rayon['nom']}_{j}"
        if st.session_state.get(stock_key, False):
            qty = st.session_state.get(stock_qty_key, 1)
            unite = st.session_state.get(stock_unit_key, "pièce")
            items.append((article, qty, unite))
    if items:
        stock_items_final[rayon["nom"]] = items

final_list = subtract_stock(final_list_before_stock, stock_items_final)

# =====================
# ONGLET 3 : MA LISTE
# =====================
with tab_liste:

    if final_list:
        st.header("📋 Ma liste de courses")

        if selected_recipes_final:
            st.caption(f"🍽️ Plats : {' • '.join(selected_recipes_final)}")

        st.divider()

        total = sum(len(items) for items in final_list.values())
        checked_count = sum(
            1
            for rayon, items in final_list.items()
            for nom, qty, unite in items
            if f"check_{rayon}_{nom}" in st.session_state.checked_items
        )
        st.progress(
            checked_count / total if total > 0 else 0,
            text=f"✅ {checked_count}/{total} articles",
        )

        for rayon, items in final_list.items():
            st.subheader(rayon)
            for nom, qty, unite in items:
                display = format_item(nom, qty, unite)
                check_key = f"check_{rayon}_{nom}"
                checked = st.checkbox(
                    display,
                    key=check_key,
                    value=check_key in st.session_state.checked_items,
                )
                if checked:
                    st.session_state.checked_items.add(check_key)
                elif check_key in st.session_state.checked_items:
                    st.session_state.checked_items.discard(check_key)

        st.divider()

        col1, col2, col3 = st.columns(3)
        with col1:
            docx_buffer = export_to_docx(final_list, selected_recipes_final)
            st.download_button(
                label="📥 Exporter en Word",
                data=docx_buffer,
                file_name=f"Liste_courses_{datetime.now().strftime('%Y-%m-%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        with col2:
            if st.button("📝 Envoyer vers Notion"):
                with st.spinner("Création de la page Notion..."):
                    success, message, url = export_to_notion(final_list, selected_recipes_final)
                if success:
                    st.success(message)
                    if url:
                        st.markdown(f"[🔗 Ouvrir dans Notion]({url})")
                else:
                    st.error(message)
        with col3:
            if st.button("🗑️ Réinitialiser les coches"):
                st.session_state.checked_items = set()
                st.rerun()
    else:
        st.info(
            "👈 Sélectionnez des recettes dans l'onglet **Recettes** "
            "ou ajoutez des articles depuis les **Produits** pour constituer votre liste."
        )
